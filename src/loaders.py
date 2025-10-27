from __future__ import annotations

import csv
from pathlib import Path
from typing import Callable, Dict, Iterable, Iterator, List, Sequence

from docx import Document as DocxDocument

from documents import Document


ChunkGenerator = Iterator[List[Document]]
LoaderFn = Callable[[Path], ChunkGenerator]


class ChunkedDocumentLoader:
    """Собирает документы с диска и выдаёт их порциями удобного размера."""

    SUPPORTED_SUFFIXES: Sequence[str] = (".txt", ".docx", ".csv", ".html", ".htm")

    def __init__(self, root: Path, chunk_size: int = 500) -> None:
        """Запоминает корень поиска и создаёт обработчики для поддерживаемых форматов."""
        self.root = Path(root)
        self.chunk_size = max(1, chunk_size)
        self._handlers: Dict[str, LoaderFn] = {
            ".txt": self._load_txt,
            ".docx": self._load_docx,
            ".csv": self._load_csv,
            ".html": self._load_html,
            ".htm": self._load_html,
        }

    def load(self) -> Iterable[List[Document]]:
        """Возвращает генератор порций документов, извлечённых из файловой системы."""
        if not self.root.exists():
            raise FileNotFoundError(f"Path {self.root} does not exist")

        paths = [self.root] if self.root.is_file() else self.root.rglob("*")
        for file_path in paths:
            handler = self._handlers.get(file_path.suffix.lower())
            if handler:
                yield from handler(file_path)

    def _load_txt(self, file_path: Path) -> ChunkGenerator:
        """Читает текстовый файл построчно и собирает непустые строки в документы."""
        chunk: List[Document] = []
        with file_path.open("r", encoding="utf-8", errors="ignore") as handle:
            for idx, line in enumerate(handle, start=1):
                text = line.strip()
                if not text:
                    continue
                chunk.append(Document(file_path, text, {"line_number": idx}))
                if len(chunk) >= self.chunk_size:
                    yield chunk
                    chunk = []
        if chunk:
            yield chunk

    def _load_docx(self, file_path: Path) -> ChunkGenerator:
        """Извлекает параграфы из DOCX и группирует их по размеру чанка."""
        chunk: List[Document] = []
        doc = DocxDocument(file_path)
        for idx, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            chunk.append(Document(file_path, text, {"paragraph": idx}))
            if len(chunk) >= self.chunk_size:
                yield chunk
                chunk = []
        if chunk:
            yield chunk

    def _load_csv(self, file_path: Path) -> ChunkGenerator:
        """Объединяет ячейки каждой строки CSV в текст и формирует документы."""
        chunk: List[Document] = []
        with file_path.open("r", encoding="utf-8", errors="ignore") as handle:
            reader = csv.reader(handle)
            header = next(reader, None)
            for idx, row in enumerate(reader, start=2):
                text = " ".join(cell.strip() for cell in row if cell.strip())
                if not text:
                    continue
                metadata = {"csv_row": idx, "header": header}
                chunk.append(Document(file_path, text, metadata))
                if len(chunk) >= self.chunk_size:
                    yield chunk
                    chunk = []
        if chunk:
            yield chunk

    def _load_html(self, file_path: Path) -> ChunkGenerator:
        """Возвращает HTML/HTM-файл единым документом без дробления."""
        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        yield [Document(file_path, raw, {"suffix": file_path.suffix.lower()})]
